import os
from docx import Document


class GetEmail:
    # Инициализируем путь и два списка
    def __init__(self):
        self.path = os.getcwd() + '\\'  # Получаем основной путь, в котором находится наш скрипт
        self.files_list = os.listdir(self.path)  # Список всех файлов в текущей папке
        # Список директорий вордов для дальнейшей обработки
        self.docx_list_dir = [self.path + docx_file for docx_file in self.files_list if '.docx' in docx_file]

    def get_emails(self):
        all_lines = list()  # создание списка, в котором будет храниться всё включая email
        for word in self.docx_list_dir:
            doc = Document(word)  # Создание переменной для работы с docx
            for para in doc.paragraphs:  # Проход по всему документу (по каждой строке)
                if '@' in para.text:  # Если в строке дока есть "@" - записываем строку разделяя её
                    all_lines.append(para.text.replace(',', '').split())  # заменяем символы

        with open('emails.txt', 'w') as file:  # Создаём блокнот, в котором будут почты со всех доков
            for i in all_lines:  # для списков в списке all_lines
                for string in i:  # для строк в списке
                    if '@' in string:  # если "@" в строке, значит это почта
                        file.write(string + '\n')  # записываем в наш блокнот почту

    def get_emails_table(self):
        all_lines = list()  # создание списка, в котором будет храниться все извлечённые строки из ворда
        for word in self.docx_list_dir:
            doc = Document(word)  # Создание переменной для работы с docx
            for table in doc.tables:  # для таблицы
                for row in table.rows:  # для ряда в рядах таблицы
                    for cell in row.cells:  # для ячейки в ряде
                        if '@' in cell.text:  # Если в тексте ячейки есть "@"
                            all_lines.append(cell.text)  # записываем строку(mail)

        with open('emails.txt', 'a+') as file:  # Добавляем почты из таблиц без перезаписи
            for mail in all_lines:  # для почт в списке all_lines(список почт)
                file.write(mail + '\n')  # записываем в наш блокнот почту

    @staticmethod
    def folders_bypass(path, docx_list_dir):
        for folder in os.listdir(path):
            if os.path.isdir(path + '\\' + folder):
                new_path = path + '\\' + folder  # Создаём переменную нового пути + имя папки
                files_list = [file for file in os.listdir(new_path) if '.docx' in file]  # список вордов
                [docx_list_dir.append(new_path + '\\' + file) for file in files_list]  # добавляем все пути с вордами
                GetEmail.folders_bypass(new_path, docx_list_dir)  # Вызываем до тех пор, пока не обойдём все папки

        return docx_list_dir


# Основная прога
def main():
    get_emails = GetEmail()
    path = get_emails.path
    docx_list_dir = get_emails.docx_list_dir
    print('Извлечение запущено ...')
    get_emails.folders_bypass(path, docx_list_dir)
    get_emails.get_emails()
    get_emails.get_emails_table()
    print('... Извлечение закончено')
    with open('emails.txt', 'r') as txt:  # Конструкция для нахождения кол-ва всех почт в нашем блокноте
        print('Всего', len(txt.readlines()), 'почт')


if __name__ == '__main__':
    main()

'''Документация
Скрипт для извлечения почт из Word.docx

        Версия скрипта: 3.1;
        Всё обёрнуто в класс;
        Скрипт рекурсивно обходит другие папки в поисках .docx;
        Скрипт работает и с таблицами и со строками, но не с гиперссылками;
        Результат сохраняется в emails.txt.
'''
